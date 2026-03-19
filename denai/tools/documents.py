"""Tools para criar documentos Word (.docx) e planilhas Excel (.xlsx)."""

from __future__ import annotations

import json
from pathlib import Path

from ..security.sandbox import is_path_allowed

# ─── Specs ─────────────────────────────────────────────────────────────────

CREATE_DOCUMENT_SPEC = {
    "type": "function",
    "function": {
        "name": "create_document",
        "description": (
            "Cria um documento Word (.docx). Suporta título, parágrafos, "
            "cabeçalhos (h1, h2, h3), listas com bullets, e tabelas. "
            "O arquivo é salvo no caminho especificado. "
            "Use para criar relatórios, cartas, contratos, documentação."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Caminho do arquivo .docx a criar "
                        "(ex: ~/Documents/relatorio.docx)"
                    ),
                },
                "content": {
                    "type": "array",
                    "description": (
                        "Lista de blocos de conteúdo. Cada bloco é um objeto com "
                        "'type' e dados. Tipos: "
                        "heading (type, text, level:1-3), "
                        "paragraph (type, text, bold:bool), "
                        "bullet_list (type, items:[str]), "
                        "table (type, headers:[str], rows:[[str]])"
                    ),
                    "items": {"type": "object"},
                },
            },
            "required": ["path", "content"],
        },
    },
}

CREATE_SPREADSHEET_SPEC = {
    "type": "function",
    "function": {
        "name": "create_spreadsheet",
        "description": (
            "Cria uma planilha Excel (.xlsx). Suporta múltiplas abas (sheets), "
            "cada uma com cabeçalhos e linhas de dados. "
            "Colunas são auto-dimensionadas. "
            "Use para criar planilhas de dados, relatórios tabulares, listas."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Caminho do arquivo .xlsx a criar "
                        "(ex: ~/Documents/dados.xlsx)"
                    ),
                },
                "sheets": {
                    "type": "array",
                    "description": (
                        "Lista de abas. Cada aba é um objeto com "
                        "'name' (nome da aba), "
                        "'headers' (lista de cabeçalhos), "
                        "'rows' (lista de listas com os dados)"
                    ),
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "headers": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                            "rows": {
                                "type": "array",
                                "items": {
                                    "type": "array",
                                    "items": {},
                                },
                            },
                        },
                        "required": ["name", "headers", "rows"],
                    },
                },
            },
            "required": ["path", "sheets"],
        },
    },
}


# ─── Executors ──────────────────────────────────────────────────────────────


def _resolve_path(raw: str) -> Path:
    """Resolve ~ e retorna Path absoluto."""
    return Path(raw).expanduser().resolve()


async def create_document(args: dict) -> str:
    """Cria um documento Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return (
            "❌ python-docx não instalado. "
            "Rode: pip install python-docx"
        )

    raw_path = args.get("path", "")
    content = args.get("content", [])

    if not raw_path:
        return "❌ Parâmetro 'path' é obrigatório."
    if not raw_path.endswith(".docx"):
        raw_path += ".docx"

    path = _resolve_path(raw_path)
    if not is_path_allowed(path):
        return f"🔒 Caminho não permitido: {path}"

    # Garante que o diretório existe
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Se content é string (JSON), faz parse
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError:
            # Texto simples — cria um parágrafo
            doc.add_paragraph(content)
            doc.save(str(path))
            return f"✅ Documento criado: {path}"

    for block in content:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            level = block.get("level", 1)
            level = max(1, min(level, 3))  # Clamp 1-3
            doc.add_heading(block.get("text", ""), level=level)

        elif block_type == "paragraph":
            p = doc.add_paragraph()
            run = p.add_run(block.get("text", ""))
            if block.get("bold"):
                run.bold = True
            if block.get("italic"):
                run.italic = True
            font_size = block.get("font_size")
            if font_size:
                run.font.size = Pt(font_size)

        elif block_type == "bullet_list":
            items = block.get("items", [])
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")

        elif block_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                # Cabeçalhos
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                # Dados
                for row_data in rows:
                    row = table.add_row()
                    for i, val in enumerate(row_data):
                        if i < len(headers):
                            row.cells[i].text = str(val)

        else:
            # Fallback: trata como parágrafo
            doc.add_paragraph(str(block.get("text", "")))

    doc.save(str(path))
    n_blocks = len(content) if isinstance(content, list) else 1
    return f"✅ Documento Word criado: {path} ({n_blocks} blocos)"


async def create_spreadsheet(args: dict) -> str:
    """Cria uma planilha Excel (.xlsx)."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
        from openpyxl.utils import get_column_letter
    except ImportError:
        return (
            "❌ openpyxl não instalado. "
            "Rode: pip install openpyxl"
        )

    raw_path = args.get("path", "")
    sheets = args.get("sheets", [])

    if not raw_path:
        return "❌ Parâmetro 'path' é obrigatório."
    if not raw_path.endswith(".xlsx"):
        raw_path += ".xlsx"

    path = _resolve_path(raw_path)
    if not is_path_allowed(path):
        return f"🔒 Caminho não permitido: {path}"

    # Garante que o diretório existe
    path.parent.mkdir(parents=True, exist_ok=True)

    # Se sheets é string (JSON), faz parse
    if isinstance(sheets, str):
        try:
            sheets = json.loads(sheets)
        except json.JSONDecodeError:
            return "❌ 'sheets' deve ser uma lista de abas."

    if not sheets:
        return "❌ Pelo menos uma aba é necessária."

    wb = Workbook()
    bold_font = Font(bold=True)

    for idx, sheet_data in enumerate(sheets):
        if idx == 0:
            ws = wb.active
            ws.title = sheet_data.get("name", "Dados")
        else:
            ws = wb.create_sheet(title=sheet_data.get("name", f"Aba{idx + 1}"))

        headers = sheet_data.get("headers", [])
        rows = sheet_data.get("rows", [])

        # Cabeçalhos em negrito
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=str(header))
            cell.font = bold_font

        # Dados
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, val in enumerate(row_data, 1):
                # Tenta converter números
                if isinstance(val, str):
                    try:
                        val = float(val) if "." in val else int(val)
                    except (ValueError, TypeError):
                        pass
                ws.cell(row=row_idx, column=col_idx, value=val)

        # Auto-dimensionar colunas
        for col_idx in range(1, len(headers) + 1):
            max_len = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 8
            for row in ws.iter_rows(
                min_row=2,
                max_row=min(len(rows) + 1, 100),  # Sample primeiras 100 rows
                min_col=col_idx,
                max_col=col_idx,
            ):
                for cell in row:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
            col_letter = get_column_letter(col_idx)
            ws.column_dimensions[col_letter].width = min(max_len + 3, 50)

    wb.save(str(path))
    total_rows = sum(len(s.get("rows", [])) for s in sheets)
    return (
        f"✅ Planilha Excel criada: {path} "
        f"({len(sheets)} aba(s), {total_rows} linhas)"
    )


# ─── Registro ──────────────────────────────────────────────────────────────

TOOLS = [
    (CREATE_DOCUMENT_SPEC, "create_document"),
    (CREATE_SPREADSHEET_SPEC, "create_spreadsheet"),
]
